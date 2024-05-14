import re
import pytesseract
from PIL import Image
from docx import Document


def analizar_imagen(i):
    try:
        nombre_imagen = "{0}. IMAGEN.jpeg".format(i + 1)
        imagen = Image.open(nombre_imagen)
        texto_extraido = pytesseract.image_to_string(imagen, lang='spa')
        return texto_extraido
    except FileNotFoundError:
        print("El archivo de imagen no fue encontrado.")
    except Exception as e:
        print("Ocurrió un error al analizar la imagen:", e)


def generar_documento_word(nombre_archivo):
    doc = Document()
    doc.save(nombre_archivo)


def guardar_texto(texto, nombre_archivo):
    doc = Document(nombre_archivo)
    doc.add_paragraph(texto)
    doc.save(nombre_archivo)


def ajustar_texto(texto):
    texto = texto.replace("-\n", "")
    texto = texto.replace("-", "")
    texto = texto.replace("\n", " ")
    texto = texto.replace("  ", " ")
    texto = texto.replace("   ", " ")
    patron_eliminar = re.compile(r'^\s*' + re.escape("Enunciado 1 17.513. ") + r'\s*', re.IGNORECASE)
    texto = patron_eliminar.sub("", texto)
    return texto


def extraer_texto_word(nombre_archivo, enunciado1, enunciado2):
    doc = Document(nombre_archivo)
    for paragraph in doc.paragraphs:
        if enunciado1 in paragraph.text:
            texto = paragraph.text
        if enunciado2 in paragraph.text:
            break
    texto = ajustar_texto(texto)
    return texto


def reconocer_tasa_media(texto):
    patron_tiempo = re.compile(
        r'(tasa media|tasa media de llegada|tiempo promedio entre llegadas|tiempo esperado entre llegadas|tasa esperada entre llegadas)\s+de\s+([1-9]|[1-9][0-9]|100)',
        re.IGNORECASE)
    patron_distribucion = re.compile(
        r'(distribución de probabilidad|distribución|distribución de)\s+(exponencial|poisson)', re.IGNORECASE)
    patron_unidades = re.compile(r'(minutos|clientes por hora|clientes por minuto)', re.IGNORECASE)

    coincidencia_tiempo = patron_tiempo.search(texto)
    coincidencia_distribucion = patron_distribucion.search(texto)
    coincidencia_unidades = patron_unidades.search(texto)

    guardar_texto("Supuesto 1", "WORD #2.docx")
    guardar_texto("Enunciado 1", "WORD #2.docx")
    guardar_texto("El nombre de la distribucion es: distribucion {0}".format(coincidencia_distribucion.group(2)),
                  "WORD #2.docx")
    guardar_texto(
        "La tasa media de llegada o tiempo promedio entre llegada es de: {}".format(coincidencia_tiempo.group(2)),
        "WORD #2.docx")
    guardar_texto("Las unidades de la tasa media de llegada o tiempo promedio entre llegada es de: {}".format(
        coincidencia_unidades.group(1)), "WORD #2.docx")
    guardar_texto(texto, "WORD #2.docx")


def reconocer_tasa_media_servicio(texto):
    patron_tiempo = re.compile(
        r'(tasa media de servicio|tasa esperada|tiempo promedio de servicio|tiempo esperado)\s+de\s+([1-9]|[1-9][0-9]|100)',
        re.IGNORECASE)
    patron_distribucion = re.compile(
        r'(distribución de probabilidad|distribución|distribución de)\s+(exponencial|poisson)', re.IGNORECASE)
    patron_unidades = re.compile(r'(minutos|clientes por hora|clientes por minuto)', re.IGNORECASE)

    coincidencia_tiempo = patron_tiempo.search(texto)
    coincidencia_distribucion = patron_distribucion.search(texto)
    coincidencia_unidades = patron_unidades.search(texto)

    guardar_texto("Supuesto 2", "WORD #2.docx")
    guardar_texto("Enunciado 1", "WORD #2.docx")
    guardar_texto("El nombre de la distribucion es: distribucion {0}".format(coincidencia_distribucion.group(2)),
                  "WORD #2.docx")
    guardar_texto(
        "La tasa media de servicio o tiempo promedio de servicio es de: {0}".format("coincidencia_tiempo.group(1)"),
        "WORD #2.docx")
    guardar_texto("Las unidades de la tasa media de servicio o tiempo promedio de servicio es de: {}".format(
        coincidencia_unidades.group(1)), "WORD #2.docx")
    guardar_texto(texto, "WORD #2.docx")


def reconocer_numero_de_servidores(texto):
    numeros_escritos = ["uno", "dos", "tres", "cuatro", "cinco", "seis", "siete", "ocho", "nueve", "diez"]
    patron_servidores = r'\b(' + '|'.join(numeros_escritos) + r')\s+servidores\b'
    coincidencia_servidores = re.search(patron_servidores, texto, re.IGNORECASE)
    guardar_texto("Supuesto 5", "WORD #2.docx")
    guardar_texto("Enunciado 1", "WORD #2.docx")
    guardar_texto("El número de servidores es: {}".format(coincidencia_servidores.group(1)), "WORD #2.docx")
    guardar_texto(texto, "WORD #2.docx")


def reconocer_capacidad_sistema(texto_extraido):
    pass


def reconocer_disciplina_atencion(texto_extraido):
    pass


def reconocer_tamano_poblacion(texto_extraido):
    pass


def main():
    print("Bienvenido al programa de análisis de imágenes.")
    generar_documento_word("WORD #1.docx")
    for i in range(3):
        texto = analizar_imagen(i)
        texto = "Enunciado {0}\n {1}".format(i + 1, texto)
        guardar_texto(texto, "WORD #1.docx")
    generar_documento_word("WORD #2.docx")
    texto_extraido = extraer_texto_word("WORD #1.docx", "Enunciado 1", "Enunciado 2")
    reconocer_tasa_media(texto_extraido)
    reconocer_tasa_media_servicio(texto_extraido)
    reconocer_numero_de_servidores(texto_extraido)
    reconocer_capacidad_sistema(texto_extraido)
    reconocer_tamano_poblacion(texto_extraido)
    reconocer_disciplina_atencion(texto_extraido)


if __name__ == "__main__":
    main()
